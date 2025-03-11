from flask import Flask, render_template, request, send_file, redirect, url_for, jsonify
from docx import Document
from docx.shared import Inches, Pt
import os
import base64
import random
import string
from werkzeug.utils import secure_filename
import imaplib
import smtplib
from email.message import EmailMessage
from num2words import num2words
from flask import send_from_directory
import json
import subprocess
import platform
from dotenv import load_dotenv
import shutil
from flask_apscheduler import APScheduler  # Nueva importación para el scheduler

# Inicializar el scheduler
scheduler = APScheduler()

def numero_a_letras(numero):
    return num2words(numero, lang='es').capitalize()

load_dotenv()

app = Flask(__name__)

if not app.debug:  # Solo en producción, no en desarrollo local
    app.config['PREFERRED_URL_SCHEME'] = 'https'

# Configuración de directorios con rutas absolutas
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_CONTRATO = os.path.join(BASE_DIR, "Plantillas", "Contrato.docx")
TEMPLATE_CERTIFICADO = os.path.join(BASE_DIR, "Plantillas", "Certificado.docx")
OUTPUT_DIR = os.path.join(BASE_DIR, "contratos_generados")
SIGNATURE_DIR = os.path.join(BASE_DIR, "firmas")
FIRMAS_TEMP_DIR = os.path.join(BASE_DIR, "firmas_temp")
os.makedirs(FIRMAS_TEMP_DIR, exist_ok=True)

# Verifica que los directorios existan
for path in [os.path.dirname(TEMPLATE_CONTRATO), OUTPUT_DIR, SIGNATURE_DIR]:
    if not os.path.exists(path):
        print(f"Creando directorio: {path}")
        os.makedirs(path, exist_ok=True)

print(f"Directorio actual: {os.getcwd()}")
print(f"Ruta de plantilla contrato: {os.path.abspath(TEMPLATE_CONTRATO)}")
print(f"Ruta de directorio de salida: {os.path.abspath(OUTPUT_DIR)}")

# Función para modificar documento
def modificar_docx(template_path, output_path, datos, firma_path=None):
    print(f"📌 Iniciando modificación de documento con plantilla: {template_path}")
    print(f"📌 Ruta de salida: {output_path}")
    
    if not os.path.exists(template_path):
        print(f"⚠️ ERROR: La plantilla no existe: {template_path}")
        raise FileNotFoundError(f"No se encontró la plantilla: {template_path}")
    
    try:
        doc = Document(template_path)
        print("✅ Plantilla cargada correctamente")
        
        def aplicar_negrita(paragraph):
            partes = paragraph.text.split("&&")
            paragraph.clear()
            for i, parte in enumerate(partes):
                run = paragraph.add_run(parte.strip())
                if i % 2 == 1:
                    run.bold = True
                run.add_text(" ")

        def wrap_text_to_width(text, width_inches, font_name='Arial', font_size=12):
            chars_per_inch = 16.67
            max_chars = int(width_inches * chars_per_inch)
            words = text.split()
            wrapped_lines = []
            current_line = []
            current_width = 0
            
            for word in words:
                word_width = len(word) * 0.06
                if current_width + word_width <= width_inches:
                    current_line.append(word)
                    current_width += word_width
                else:
                    if current_line:
                        wrapped_lines.append(" ".join(current_line))
                    current_line = [word]
                    current_width = word_width
            
            if current_line:
                wrapped_lines.append(" ".join(current_line))
            return "\n".join(wrapped_lines)

        processed = set()
        for p in doc.paragraphs:
            if p.text.strip() and "{{FIRMAS_TEXTO}}" in p.text.strip() and id(p) not in processed:
                processed.add(id(p))
                firmas_texto = datos.get("{{FIRMAS_TEXTO}}", "").strip()
                firma_paths = datos.get("firma_paths", [])
                print(f"Encontrado {{FIRMAS_TEXTO}} en párrafo: {p.text}")
                print(f"Contenido de firmas_texto: {firmas_texto}")
                
                if firmas_texto:
                    p.clear()
                    
                    if "\n" in firmas_texto and firmas_texto.count("\n") < 4:
                        lines = firmas_texto.split("\n")
                        if lines:
                            if "___________________________" in lines[0] and firma_paths:
                                run = p.add_run()
                                run.add_picture(firma_paths[0], width=Inches(1.6))
                                p.add_run("\n___________________________")
                                p.add_run("\n" + "\n".join(lines[1:]))
                            else:
                                p.text = "\n".join(lines)
                            p.alignment = 0
                            for run in p.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(12)
                                run.bold = True
                    
                    elif firmas_texto.startswith("COL1|COL2"):
                        table_data = firmas_texto.replace("COL1|COL2", "").strip().split("|")
                        if len(table_data) == 2:
                            p.clear()
                            table = doc.add_table(rows=1, cols=2)
                            table.alignment = 1
                            table.autofit = False
                            table.width = Inches(6.14)
                            
                            for col in table.columns:
                                col.width = Inches(2.82)
                                for cell in col.cells:
                                    for para in cell.paragraphs:
                                        para.alignment = 1
                            
                            col1_text = table_data[0].strip().split("\n")
                            col2_text = table_data[1].strip().split("\n")
                            
                            cell1 = table.rows[0].cells[0]
                            cell2 = table.rows[0].cells[1]
                            
                            para1 = cell1.add_paragraph()
                            if col1_text:
                                if "___________________________" in col1_text[0] and len(firma_paths) > 0:
                                    run = para1.add_run()
                                    run.add_picture(firma_paths[0], width=Inches(1.6))
                                    para1.add_run("\n___________________________")
                                    para1.add_run("\n" + "\n".join(col1_text[1:]))
                                else:
                                    para1.text = "\n".join(col1_text)
                            
                            para2 = cell2.add_paragraph()
                            if col2_text:
                                if "___________________________" in col2_text[0] and len(firma_paths) > 1:
                                    run = para2.add_run()
                                    run.add_picture(firma_paths[1], width=Inches(1.6))
                                    para2.add_run("\n___________________________")
                                    para2.add_run("\n" + "\n".join(col2_text[1:]))
                                else:
                                    para2.text = "\n".join(col2_text)
                            
                            for cell in [cell1, cell2]:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.font.name = 'Arial'
                                        run.font.size = Pt(12)
                                        run.bold = True
                            p._element.getparent().replace(p._element, table._element)
                    
                    elif "\n\n" in firmas_texto:
                        lines = firmas_texto.split("\n\n")
                        for i, line_block in enumerate(lines):
                            p.clear()
                            block_lines = line_block.strip().split("\n")
                            if block_lines:
                                if "___________________________" in block_lines[0] and len(firma_paths) > i:
                                    run = p.add_run()
                                    run.add_picture(firma_paths[i], width=Inches(1.6))
                                    p.add_run("\n___________________________")
                                    p.add_run("\n" + "\n".join(block_lines[1:]))
                                else:
                                    p.text = "\n".join(block_lines)
                                p.alignment = 1
                                for run in p.runs:
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(12)
                                    run.bold = True
                            if i < len(lines) - 1:
                                doc.add_paragraph("")
            else:
                for key, value in datos.items():
                    if key in p.text and key != "{{FIRMAS_TEXTO}}":
                        for run in p.runs:
                            if key in run.text:
                                run.text = run.text.replace(key, value)
                if "&&" in p.text:
                    aplicar_negrita(p)
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
        
        if firma_path:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if "{{FIRMA}}" in para.text:
                                para.text = ""
                                run = para.add_run()
                                run.add_picture(firma_path, width=Inches(1.6))

        doc.save(output_path)
        print(f"✅ Documento guardado correctamente en: {output_path}")
        return True
    except Exception as e:
        print(f"⚠️ ERROR en modificar_docx: {str(e)}")
        raise e

def wrap_text_to_width(text, width_inches, font_name='Arial', font_size=12):
    chars_per_inch = 16.67
    max_chars = int(width_inches * chars_per_inch)
    words = text.split()
    wrapped_lines = []
    current_line = []
    current_width = 0
    
    for word in words:
        word_width = len(word) * 0.06
        if current_width + word_width <= width_inches:
            current_line.append(word)
            current_width += word_width
        else:
            if current_line:
                wrapped_lines.append(" ".join(current_line))
            current_line = [word]
            current_width = word_width
    
    if current_line:
        wrapped_lines.append(" ".join(current_line))
    return "\n".join(wrapped_lines)

def convertir_docx_a_pdf(docx_path, pdf_path):
    try:
        if platform.system() == "Windows":
            libreoffice_cmd = r"C:\Program Files\LibreOffice\program\soffice.exe"
        else:
            libreoffice_cmd = shutil.which("libreoffice")
            
        print(f"🔍 Ruta de LibreOffice: {libreoffice_cmd}")
        if libreoffice_cmd is None:
            raise FileNotFoundError("❌ LibreOffice no está instalado o no está en el PATH")

        subprocess.run(
            [libreoffice_cmd, "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(pdf_path), docx_path],
            check=True
        )
        return os.path.exists(pdf_path)
    except Exception as e:
        print(f"❌ Error en la conversión: {e}")
        return False

@app.route('/contratos_generados/<filename>')
def ver_contrato(filename):
    return send_from_directory(OUTPUT_DIR, filename)

def guardar_en_enviados(msg, email_user_gmail, email_pass_gmail, imap_server="imap.gmail.com"):
    SENT_FOLDER = '"[Gmail]/Enviados"'
    try:
        mail = imaplib.IMAP4_SSL(imap_server)
        mail.login(email_user_gmail, email_pass_gmail)
        mail.select(SENT_FOLDER)
        raw_message = msg.as_bytes()
        mail.append(SENT_FOLDER, None, None, raw_message)
        mail.logout()
        print("✅ Correo guardado en 'Enviados' de Gmail")
        return True
    except Exception as e:
        print(f"❌ Error al guardar en 'Enviados': {str(e)} (Usuario: {email_user_gmail})")
        return False

def enviar_correo(destinatario, asunto, cuerpo, adjunto, guardar_en_gmail=True):
    try:
        msg = EmailMessage()
        msg['Subject'] = asunto
        msg['From'] = os.getenv("EMAIL_USER_HOSTINGER", "asesorias@grupotrujilloyasociados.com")
        
        if isinstance(destinatario, (list, tuple)):
            msg['To'] = destinatario[0]
            if len(destinatario) > 1:
                msg['Cc'] = destinatario[1]
        else:
            msg['To'] = destinatario
            
        msg.set_content(cuerpo)
        
        if not os.path.exists(adjunto):
            print(f"⚠️ ERROR: El archivo adjunto no existe: {adjunto}")
            return False
            
        with open(adjunto, 'rb') as f:
            file_data = f.read()
            file_name = os.path.basename(adjunto)
            msg.add_attachment(file_data, maintype='application', subtype='pdf', filename=file_name)
        
        with smtplib.SMTP('smtp.hostinger.com', 587) as server:
            server.starttls()
            email_user_hostinger = os.getenv("EMAIL_USER_HOSTINGER", "asesorias@grupotrujilloyasociados.com")
            email_pass_hostinger = os.getenv("EMAIL_PASS_HOSTINGER")
            if not email_pass_hostinger:
                raise ValueError("La contraseña de Hostinger no está configurada en las variables de entorno")
            server.login(email_user_hostinger, email_pass_hostinger)
            server.send_message(msg)
            print(f"✅ Correo enviado correctamente a {destinatario}")
            
            if guardar_en_gmail:
                email_user_gmail = os.getenv("EMAIL_USER_GMAIL", "grupojuridicotrujillo@gmail.com")
                email_pass_gmail = os.getenv("EMAIL_PASS_GMAIL")
                if not email_pass_gmail:
                    raise ValueError("La contraseña de Gmail no está configurada en las variables de entorno")
                guardar_en_enviados(msg, email_user_gmail, email_pass_gmail)
            return True
    except Exception as e:
        print(f"❌ Error al enviar correo: {e}")
        return False

@app.route('/convertir_monto')
def convertir_monto():
    valor = request.args.get('valor', default="0").replace(".", "")
    if valor.isdigit():
        return numero_a_letras(int(valor))
    return "Número inválido"

@app.route('/vista_previa_y_firmar')
def vista_previa_y_firmar():
    if not os.path.exists(TEMPLATE_CONTRATO):
        print(f"⚠️ ERROR: La plantilla no existe: {TEMPLATE_CONTRATO}")
        return f"Error: No se encontró la plantilla en {TEMPLATE_CONTRATO}", 500

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    test_file = os.path.join(OUTPUT_DIR, "test_write.txt")
    try:
        with open(test_file, 'w') as f:
            f.write("Test")
        os.remove(test_file)
    except Exception as e:
        print(f"⚠️ ERROR: No se puede escribir en el directorio {OUTPUT_DIR}: {str(e)}")
        return f"Error: No hay permisos de escritura en {OUTPUT_DIR}", 500

    mandantes = request.args.getlist('mandantes[]')
    cedulas = request.args.getlist('cedulas[]')
    dni_mandantes = request.args.getlist('tipo_identificacion_mandantes[]')
    demandados = request.args.getlist('demandados[]')
    cedula_demandados = request.args.getlist('nit_cc_demandados[]')
    dni_demandados = request.args.getlist('tipo_identificacion_demandados[]')
    pago_final = request.args.get('pago_final', 'No')
    porcentaje = request.args.get('porcentaje', '')
    email = request.args.get('email', '')
    numero_contrato = request.args.get('numero_contrato', '')
    monto = request.args.get('monto', '0')
    monto_letras = request.args.get('monto_letras', numero_a_letras(int(monto.replace(".", ""))).upper())

    if pago_final == 'Si':
        texto_pago_final = (
            f", si el proceso termina de manera favorable en sentencia judicial o por medio de conciliación, "
            f"se cancelará adicionalmente por parte del MANDANTE la suma del {porcentaje}% sobre las pretensiones obtenidas en la demanda"
        )
    else:
        texto_pago_final = ""

    if mandantes and cedulas and dni_mandantes and len(mandantes) == len(cedulas) == len(dni_mandantes):
        if len(mandantes) == 1:
            firmas_texto = f"___________________________\n{mandantes[0]}\n{dni_mandantes[0]} {cedulas[0]}"
        elif len(mandantes) == 2:
            firmas_texto = (
                "COL1|COL2"
                "___________________________"
                f"\n{mandantes[0]}"
                f"\n{dni_mandantes[0]} {cedulas[0]}"
                "|___________________________"
                f"\n{mandantes[1]}"
                f"\n{dni_mandantes[1]} {cedulas[1]}"
            )
        else:
            firmas_texto = "\n".join([f"___________________________\\n{mandantes[i]}\\n{dni_mandantes[i]} {cedulas[i]}" for i in range(len(mandantes))])
    else:
        firmas_texto = ""

    if mandantes and cedulas and dni_mandantes and len(mandantes) == len(cedulas) == len(dni_mandantes):
        mandantes_texto = f"&& {mandantes[0]} &&, mayor y vecino/a de esta ciudad, identificado/a con {dni_mandantes[0]} N° && {cedulas[0]} && "
        for i in range(1, len(mandantes)):
            mandantes_texto += f" y && {mandantes[i]} &&, mayor y vecino/a de esta ciudad, identificado/a con {dni_mandantes[i]} N° && {cedulas[i]} &&"
    else:
        mandantes_texto = "No se especificaron mandantes."

    if demandados and cedula_demandados and dni_demandados and len(demandados) == len(cedula_demandados) == len(dni_demandados):
        demandado_texto = f"&& {demandados[0]} &&, identificado/a con {dni_demandados[0]} N° && {cedula_demandados[0]} &&"
        for i in range(1, len(demandados)):
            demandado_texto += f" y && {demandados[i]} &&, identificado/a con {dni_demandados[i]} N° && {cedula_demandados[i]} &&"
    else:
        demandado_texto = "No se encontraron demandados"

    datos = {
        "{{NUMERO_CONTRATO}}": numero_contrato,
        "{{MANDANTES}}": mandantes_texto,
        "{{FIRMAS_TEXTO}}": firmas_texto,
        "{{PAGO_FINAL}}": texto_pago_final,
        "{{PORCENTAJE}}": porcentaje,
        "{{DEMANDADOS}}": demandado_texto,
        "{{MONTO}}": monto,
        "{{MONTO_LETRAS}}": monto_letras
    }

    nombre_base = secure_filename(numero_contrato.upper()) if numero_contrato else "sin_numero_contrato"
    contrato_docx = os.path.join(OUTPUT_DIR, f"contrato_{nombre_base}.docx")
    contrato_pdf = os.path.join(OUTPUT_DIR, f"contrato_{nombre_base}.pdf")

    try:
        modificar_docx(TEMPLATE_CONTRATO, contrato_docx, datos)
        print(f"✅ Documento Word creado: {contrato_docx}")

        if not os.path.exists(contrato_docx):
            return f"Error: No se pudo crear el archivo Word. Verifique permisos y rutas.", 500

        convertir_docx_a_pdf(contrato_docx, contrato_pdf)

        if not os.path.exists(contrato_pdf):
            return f"Error: No se pudo convertir el documento a PDF.", 500

        firma_params = {
            'mandantes[]': mandantes,
            'cedulas[]': cedulas,
            'tipo_identificacion_mandantes[]': dni_mandantes,
            'demandados[]': demandados,
            'nit_cc_demandados[]': cedula_demandados,
            'tipo_identificacion_demandados[]': dni_demandados,
            'numero_contrato': numero_contrato,
            'monto': monto,
            'monto_letras': monto_letras,
            'pago_final': pago_final,
            'porcentaje': porcentaje,
            'email': email
        }

        enlace_firma = url_for('firma', nombre_archivo=nombre_base, **firma_params)
        print(f"Redirigiendo a: {enlace_firma}")
        return redirect(enlace_firma)

    except Exception as e:
        print(f"Error en vista_previa_y_firmar: {str(e)}")
        return f"Error generando vista previa y redirigiendo: {str(e)}", 500
    
@app.route('/vista_previa', methods=['GET'])
def vista_previa():
    # Verificar que la plantilla exista
    if not os.path.exists(TEMPLATE_CONTRATO):
        print(f"⚠️ ERROR: La plantilla no existe: {TEMPLATE_CONTRATO}")
        return f"Error: No se encontró la plantilla en {TEMPLATE_CONTRATO}", 500

    # Crear el directorio de salida si no existe
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Obtener los parámetros del formulario enviados por GET
    mandantes = request.args.getlist('mandantes[]')
    cedulas = request.args.getlist('cedulas[]')
    dni_mandantes = request.args.getlist('tipo_identificacion_mandantes[]')
    demandados = request.args.getlist('demandados[]')
    cedula_demandados = request.args.getlist('nit_cc_demandados[]')
    dni_demandados = request.args.getlist('tipo_identificacion_demandados[]')
    pago_final = request.args.get('pago_final', 'No')
    porcentaje = request.args.get('porcentaje', '')
    email = request.args.get('email', '')
    numero_contrato = request.args.get('numero_contrato', '')
    monto = request.args.get('monto', '0')
    monto_letras = request.args.get('monto_letras', numero_a_letras(int(monto.replace(".", ""))).upper())

    # Generar texto de pago final
    if pago_final == 'Si':
        texto_pago_final = (
            f", si el proceso termina de manera favorable en sentencia judicial o por medio de conciliación, "
            f"se cancelará adicionalmente por parte del MANDANTE la suma del {porcentaje}% sobre las pretensiones obtenidas en la demanda"
        )
    else:
        texto_pago_final = ""

    # Generar texto de firmas
    if mandantes and cedulas and dni_mandantes and len(mandantes) == len(cedulas) == len(dni_mandantes):
        if len(mandantes) == 1:
            firmas_texto = f"___________________________\n{mandantes[0]}\n{dni_mandantes[0]} {cedulas[0]}"
        elif len(mandantes) == 2:
            firmas_texto = (
                "COL1|COL2"
                "___________________________"
                f"\n{mandantes[0]}"
                f"\n{dni_mandantes[0]} {cedulas[0]}"
                "|___________________________"
                f"\n{mandantes[1]}"
                f"\n{dni_mandantes[1]} {cedulas[1]}"
            )
        else:
            firmas_texto = "\n".join([f"___________________________\\n{mandantes[i]}\\n{dni_mandantes[i]} {cedulas[i]}" for i in range(len(mandantes))])
    else:
        firmas_texto = ""

    # Generar texto de mandantes
    if mandantes and cedulas and dni_mandantes and len(mandantes) == len(cedulas) == len(dni_mandantes):
        mandantes_texto = f"&& {mandantes[0]} &&, mayor y vecino/a de esta ciudad, identificado/a con {dni_mandantes[0]} N° && {cedulas[0]} && "
        for i in range(1, len(mandantes)):
            mandantes_texto += f" y && {mandantes[i]} &&, mayor y vecino/a de esta ciudad, identificado/a con {dni_mandantes[i]} N° && {cedulas[i]} &&"
    else:
        mandantes_texto = "No se especificaron mandantes."

    # Generar texto de demandados
    if demandados and cedula_demandados and dni_demandados and len(demandados) == len(cedula_demandados) == len(dni_demandados):
        demandado_texto = f"&& {demandados[0]} &&, identificado/a con {dni_demandados[0]} N° && {cedula_demandados[0]} &&"
        for i in range(1, len(demandados)):
            demandado_texto += f" y && {demandados[i]} &&, identificado/a con {dni_demandados[i]} N° && {cedula_demandados[i]} &&"
    else:
        demandado_texto = "No se encontraron demandados"

    # Preparar los datos para el documento preliminar
    datos = {
        "{{NUMERO_CONTRATO}}": numero_contrato,
        "{{MANDANTES}}": mandantes_texto,
        "{{FIRMAS_TEXTO}}": firmas_texto,
        "{{PAGO_FINAL}}": texto_pago_final,
        "{{PORCENTAJE}}": porcentaje,
        "{{DEMANDADOS}}": demandado_texto,
        "{{MONTO}}": monto,
        "{{MONTO_LETRAS}}": monto_letras
    }

    # Generar nombre del archivo
    nombre_base = secure_filename(numero_contrato.upper()) if numero_contrato else "vista_previa_temp"
    contrato_docx = os.path.join(OUTPUT_DIR, f"contrato_{nombre_base}.docx")
    contrato_pdf = os.path.join(OUTPUT_DIR, f"contrato_{nombre_base}.pdf")

    try:
        # Generar el documento Word preliminar
        modificar_docx(TEMPLATE_CONTRATO, contrato_docx, datos)
        print(f"✅ Documento Word creado: {contrato_docx}")

        if not os.path.exists(contrato_docx):
            return f"Error: No se pudo crear el archivo Word. Verifique permisos y rutas.", 500

        # Convertir a PDF
        convertir_docx_a_pdf(contrato_docx, contrato_pdf)

        if not os.path.exists(contrato_pdf):
            return f"Error: No se pudo convertir el documento a PDF.", 500

        # Generar la URL del PDF para el iframe
        pdf_url = url_for('ver_contrato', filename=f"contrato_{nombre_base}.pdf", _external=True, _scheme='https')

        # Renderizar la plantilla de vista previa con la URL del PDF
        return render_template('vista_previa.html', pdf_url=pdf_url)

    except Exception as e:
        print(f"Error en vista_previa: {str(e)}")
        return f"Error generando vista previa: {str(e)}", 500

@app.route('/', methods=['GET', 'POST'])
def formulario():
    if request.method == 'POST':
        mandantes = request.form.getlist('mandantes[]')
        cedulas = request.form.getlist('cedulas[]')
        dni_mandantes = request.form.getlist('tipo_identificacion_mandantes[]')
        demandados = request.form.getlist('demandados[]')
        cedula_demandados = request.form.getlist('nit_cc_demandados[]')
        dni_demandados = request.form.getlist('tipo_identificacion_demandados[]')
        pago_final = request.form.get('pago_final', 'No')
        porcentaje = request.form.get('porcentaje', '')
        email = request.form.get('email', '')
        numero_contrato = request.form.get('numero_contrato', '')
        monto_str = request.form.get('monto', '0').replace(".", "")
        monto = int(monto_str) if monto_str.isdigit() else 0
        monto_letras = numero_a_letras(monto).upper()

        if pago_final == 'Si':
            texto_pago_final = (
                f", si el proceso termina de manera favorable en sentencia judicial o por medio de conciliación, "
                f"se cancelará adicionalmente por parte del MANDANTE la suma del {porcentaje}% sobre las pretensiones obtenidas en la demanda"
            )
        else:
            texto_pago_final = ""

        if mandantes and cedulas and dni_mandantes and len(mandantes) == len(cedulas) == len(dni_mandantes):
            if len(mandantes) == 1:
                firmas_texto = f"___________________________\n{mandantes[0]}\n{dni_mandantes[0]} {cedulas[0]}"
            elif len(mandantes) == 2:
                firmas_texto = (
                    "COL1|COL2"
                    "___________________________"
                    f"\n{mandantes[0]}"
                    f"\n{dni_mandantes[0]} {cedulas[0]}"
                    "|___________________________"
                    f"\n{mandantes[1]}"
                    f"\n{dni_mandantes[1]} {cedulas[1]}"
                )
            else:
                firmas_texto = "\n".join([f"___________________________\\n{mandantes[i]}\\n{dni_mandantes[i]} {cedulas[i]}" for i in range(len(mandantes))])
        else:
            firmas_texto = ""

        if mandantes and cedulas and dni_mandantes and len(mandantes) == len(cedulas) == len(dni_mandantes):
            mandantes_texto = f"&& {mandantes[0]} &&, mayor y vecino/a de esta ciudad, identificado/a con {dni_mandantes[0]} N° && {cedulas[0]} && "
            for i in range(1, len(mandantes)):
                mandantes_texto += f" y && {mandantes[i]} &&, mayor y vecino/a de esta ciudad, identificado/a con {dni_mandantes[i]} N° && {cedulas[i]} &&"
        else:
            mandantes_texto = "No se especificaron mandantes."

        if demandados and cedula_demandados and dni_demandados and len(demandados) == len(cedula_demandados) == len(dni_demandados):
            demandado_texto = f"&& {demandados[0]} &&, identificado/a con {dni_demandados[0]} N° && {cedula_demandados[0]} &&"
            for i in range(1, len(demandados)):
                demandado_texto += f" y && {demandados[i]} &&, identificado/a con {dni_demandados[i]} N° && {cedula_demandados[i]} &&"
        else:
            demandado_texto = "No se encontraron demandados"

        datos = {
            "{{NUMERO_CONTRATO}}": numero_contrato,
            "{{PAGO_FINAL}}": texto_pago_final,
            "{{MANDANTES}}": mandantes_texto,
            "{{FIRMAS_TEXTO}}": firmas_texto,
            "{{DEMANDADOS}}": demandado_texto,
            "{{PORCENTAJE}}": porcentaje,
            "{{MONTO}}": request.form['monto'],
            "{{MONTO_LETRAS}}": monto_letras
        }

        if mandantes and len(mandantes) > 0:
            nombre_archivo = " Y ".join(mandantes).replace(" ", "_").replace("/", "_").replace("\\", "_")
        else:
            nombre_archivo = "sin_mandantes"

        contrato_docx = os.path.join(OUTPUT_DIR, f"{nombre_archivo}.docx")
        contrato_pdf = os.path.join(OUTPUT_DIR, f"{nombre_archivo}.pdf")

        try:
            modificar_docx(TEMPLATE_CONTRATO, contrato_docx, datos)
            convertir_docx_a_pdf(contrato_docx, contrato_pdf)

            if not os.path.exists(contrato_pdf):
                return jsonify({
                    "success": False,
                    "message": "Error: No se pudo generar el PDF del contrato."
                }), 500

            firma_params = {
                'nombre_archivo': nombre_archivo,
                'mandantes[]': mandantes,
                'cedulas[]': cedulas,
                'tipo_identificacion_mandantes[]': dni_mandantes,
                'demandados[]': demandados,
                'nit_cc_demandados[]': cedula_demandados,
                'tipo_identificacion_demandados[]': dni_demandados,
                'numero_contrato': numero_contrato,
                'monto': request.form['monto'],
                'monto_letras': monto_letras,
                'pago_final': pago_final,
                'porcentaje': porcentaje,
                'email': email
            }
            enlace_firma = url_for('firma', _external=True, **firma_params)

            cuerpo_correo = (
                "Cordial saludo,\n\n"
                "Adjuntamos el documento de autorización para su revisión. Para proceder con la firma del contrato, "
                "por favor haga clic en el siguiente enlace:\n\n"
                f"{enlace_firma}\n\n"
                "Si tiene alguna duda, no dude en contactarnos.\n\n"
                "Atentamente,\nGrupo Trujillo y Asociados"
            )

            enviar_correo(
                email,
                "Trujillo y Asociados - Enlace para Firmar Contrato",
                cuerpo_correo,
                contrato_pdf
            )

            return jsonify({
                "success": True,
                "message": f"Se ha enviado un correo a {email} con el enlace para firmar el contrato.",
                "enlace_firma": enlace_firma
            })

        except Exception as e:
            print(f"Error en la generación del contrato o envío de correo: {str(e)}")
            return jsonify({
                "success": False,
                "message": f"Error al generar el contrato o enviar el correo: {str(e)}"
            }), 500

    return render_template('formulario.html')

def guardar_firma_temporal(nombre_archivo, mandante, firma_data):
    firma_path = os.path.join(FIRMAS_TEMP_DIR, f"{nombre_archivo}_{mandante}.png")
    estado_path = os.path.join(FIRMAS_TEMP_DIR, f"{nombre_archivo}_estado.json")
    
    with open(firma_path, "wb") as fh:
        fh.write(base64.b64decode(firma_data.split(",")[1]))
    
    estado = {}
    if os.path.exists(estado_path):
        with open(estado_path, "r") as f:
            estado = json.load(f)
    estado[mandante] = f"{nombre_archivo}_{mandante}.png"
    with open(estado_path, "w") as f:
        json.dump(estado, f)
    return firma_path

def cargar_firmas_existentes(nombre_archivo):
    estado_path = os.path.join(FIRMAS_TEMP_DIR, f"{nombre_archivo}_estado.json")
    if os.path.exists(estado_path):
        with open(estado_path, "r") as f:
            estado = json.load(f)
        firmas = {}
        for mandante, firma_nombre in estado.items():
            firma_path = os.path.join(FIRMAS_TEMP_DIR, firma_nombre)
            if os.path.exists(firma_path):
                with open(firma_path, "rb") as fh:
                    firmas[mandante] = "data:image/png;base64," + base64.b64encode(fh.read()).decode("utf-8")
        return firmas
    return {}

@app.route('/guardar_firma_temporal/<nombre_archivo>', methods=['POST'])
def guardar_firma_temporal_route(nombre_archivo):
    mandante = request.form.get('mandante')
    firma_data = request.form.get('firma_data')
    if not mandante or not firma_data:
        return jsonify({"mensaje": "Faltan datos (mandante o firma)."}), 400
    
    try:
        firma_path = guardar_firma_temporal(nombre_archivo, mandante, firma_data)
        return jsonify({"mensaje": f"Firma guardada exitosamente para {mandante}", "firma_path": firma_path})
    except Exception as e:
        return jsonify({"mensaje": f"Error al guardar la firma: {str(e)}"}), 500

@app.route('/firma/<nombre_archivo>', methods=['GET', 'POST'])
def firma(nombre_archivo):
    if request.method == 'POST':
        firmas_data = request.form.getlist('firma_data[]')
        
        if not firmas_data or all(not firma.strip() for firma in firmas_data):
            return jsonify({"mensaje": "Error: No se recibieron firmas válidas"}), 400

        firma_paths = []
        for i, firma_data in enumerate(firmas_data):
            if firma_data:
                try:
                    firma_path = os.path.join(SIGNATURE_DIR, f"firma_{nombre_archivo}_{i+1}.png")
                    with open(firma_path, "wb") as fh:
                        fh.write(base64.b64decode(firma_data.split(",")[1]))
                    firma_paths.append(firma_path)
                except Exception as e:
                    print(f"Error al guardar firma {i+1}: {str(e)}")
                    return jsonify({"mensaje": f"Error al guardar la firma {i+1}: {str(e)}"}), 500
                
        mandantes = request.form.getlist('mandantes[]') or request.args.getlist('mandantes[]')
        cedulas = request.form.getlist('cedulas[]') or request.args.getlist('cedulas[]')
        dni_mandantes = request.form.getlist('tipo_identificacion_mandantes[]') or request.args.getlist('tipo_identificacion_mandantes[]')
        demandados = request.form.getlist('demandados[]') or request.args.getlist('demandados[]')
        cedula_demandados = request.form.getlist('nit_cc_demandados[]') or request.args.getlist('nit_cc_demandados[]')
        dni_demandados = request.form.getlist('tipo_identificacion_demandados[]') or request.args.getlist('tipo_identificacion_demandados[]')
        numero_contrato = request.form.get('numero_contrato', request.args.get('numero_contrato', nombre_archivo))
        monto = request.form.get('monto', request.args.get('monto', '0'))
        monto_letras = request.form.get('monto_letras', request.args.get('monto_letras', numero_a_letras(int(monto.replace(".", ""))).upper()))
        pago_final = request.form.get('pago_final', request.args.get('pago_final', 'No'))
        porcentaje = request.form.get('porcentaje', request.args.get('porcentaje', ''))
        email = request.form.get('email', request.args.get('email', ''))

        if not mandantes:
            mandantes = ["Mandante_no_especificado"]
        if len(cedulas) < len(mandantes):
            cedulas.extend(["N/A"] * (len(mandantes) - len(cedulas)))
        if len(dni_mandantes) < len(mandantes):
            dni_mandantes.extend(["C.C."] * (len(mandantes) - len(dni_mandantes)))

        texto_pago_final = (
            f", si el proceso termina de manera favorable en sentencia judicial o por medio de conciliación, "
            f"se cancelará adicionalmente por parte del MANDANTE la suma del {porcentaje}% sobre las pretensiones obtenidas en la demanda"
            if pago_final == 'Si' else ""
        )

        if mandantes and len(mandantes) == len(cedulas) == len(dni_mandantes):
            mandantes_texto = f"&& {mandantes[0]} &&, mayor y vecino/a de esta ciudad, identificado/a con {dni_mandantes[0]} N° && {cedulas[0]} && "
            for i in range(1, len(mandantes)):
                mandantes_texto += f" y && {mandantes[i]} &&, mayor y vecino/a de esta ciudad, identificado/a con {dni_mandantes[i]} N° && {cedulas[i]} &&"
        else:
            mandantes_texto = "No se especificaron mandantes."

        if demandados and len(demandados) == len(cedula_demandados) == len(dni_demandados):
            demandado_texto = f"&& {demandados[0]} &&, identificado/a con {dni_demandados[0]} N° && {cedula_demandados[0]} &&"
            for i in range(1, len(demandados)):
                demandado_texto += f" y && {demandados[i]} &&, identificado/a con {dni_demandados[i]} N° && {cedula_demandados[i]} &&"
        else:
            demandado_texto = "No se encontraron demandados"

        if len(mandantes) == 1:
            firmas_texto = f"___________________________\n{mandantes[0]}\n{dni_mandantes[0]} {cedulas[0]}"
        elif len(mandantes) == 2:
            firmas_texto = (
                "COL1|COL2\n"
                "___________________________\n" +
                mandantes[0] + "\n" +
                f"{dni_mandantes[0]} {cedulas[0]}" +
                "\n|___________________________\n" +
                mandantes[1] + "\n" +
                f"{dni_mandantes[1]} {cedulas[1]}"
            )
        else:
            firmas_texto = "\n".join([f"___________________________\\n{mandantes[i]}\\n{dni_mandantes[i]} {cedulas[i]}" for i in range(len(mandantes))])

        datos = {
            "{{NUMERO_CONTRATO}}": numero_contrato,
            "{{MANDANTES}}": mandantes_texto,
            "{{FIRMAS_TEXTO}}": firmas_texto,
            "{{DEMANDADOS}}": demandado_texto,
            "{{MONTO}}": monto,
            "{{MONTO_LETRAS}}": monto_letras,
            "{{PAGO_FINAL}}": texto_pago_final,
            "{{PORCENTAJE}}": porcentaje,
            "firma_paths": firma_paths
        }

        nombre_base = secure_filename(" Y ".join(mandantes).replace(" ", "_").replace("/", "_").replace("\\", "_"))
        contrato_firmado_docx = os.path.join(OUTPUT_DIR, f"contrato_{nombre_base}.docx")
        contrato_firmado_pdf = os.path.join(OUTPUT_DIR, f"contrato_{nombre_base}.pdf")

        try:
            modificar_docx(TEMPLATE_CONTRATO, contrato_firmado_docx, datos)
            convertir_docx_a_pdf(contrato_firmado_docx, contrato_firmado_pdf)

            if not os.path.exists(contrato_firmado_pdf):
                return jsonify({"mensaje": "Error: No se pudo generar el documento firmado"}), 500

            if email:
                print(f"Enviando correo al cliente: {email}")
                cuerpo_cliente = (
                    "Cordial saludo,\n\n"
                    "Adjunto encontrarás tu contrato firmado. Gracias por confiar en nosotros para tu trámite legal.\n\n"
                    "Atentamente,\nGrupo Trujillo y Asociados"
                )
                enviar_correo(
                    email,
                    "Contrato Firmado - Grupo Trujillo y Asociados",
                    cuerpo_cliente,
                    contrato_firmado_pdf
                )
            else:
                print("⚠️ No se proporcionó un email para el cliente desde el formulario inicial.")

            print(f"Enviando copia al correo empresarial: asesorias@grupotrujilloyasociados.com")
            cuerpo_empresa = (
                f"Se ha firmado el contrato {nombre_base}.\n\n"
                "Adjunto se encuentra el documento firmado para su registro.\n\n"
                "Saludos,\nSistema Automático"
            )
            enviar_correo(
                "asesorias@grupotrujilloyasociados.com",
                f"Contrato Firmado: {nombre_base}",
                cuerpo_empresa,
                contrato_firmado_pdf
            )

            return jsonify({
                "mensaje": "Contrato firmado generado y enviado con éxito. Gracias por confiar en nosotros.",
                "redirect": url_for('descargar', tipo='contrato', numero_contrato=nombre_base)
            })

        except Exception as e:
            print(f"Error al generar contrato firmado o enviar correos: {str(e)}")
            return jsonify({"mensaje": f"Error al procesar el documento o enviar correos: {str(e)}"}), 500

    mandantes = request.args.getlist('mandantes[]')
    cedulas = request.args.getlist('cedulas[]')
    dni_mandantes = request.args.getlist('tipo_identificacion_mandantes[]')
    demandados = request.args.getlist('demandados[]')
    cedula_demandados = request.args.getlist('nit_cc_demandados[]')
    dni_demandados = request.args.getlist('tipo_identificacion_demandados[]')
    numero_contrato = request.args.get('numero_contrato', nombre_archivo)
    monto = request.args.get('monto', '0')
    monto_letras = request.args.get('monto_letras', numero_a_letras(int(monto.replace(".", ""))).upper())
    pago_final = request.args.get('pago_final', 'No')
    porcentaje = request.args.get('porcentaje', '')
    email = request.args.get('email', '')

    if not mandantes:
        mandantes = ["Mandante_no_especificado"]

    firmas_existentes = cargar_firmas_existentes(nombre_archivo)

    return render_template(
        'firma.html',
        nombre_archivo=nombre_archivo,
        mandantes=mandantes,
        cedulas=cedulas,
        tipo_identificacion_mandantes=dni_mandantes,
        demandados=demandados,
        nit_cc_demandados=cedula_demandados,
        tipo_identificacion_demandados=dni_demandados,
        numero_contrato=numero_contrato,
        monto=monto,
        monto_letras=monto_letras,
        pago_final=pago_final,
        porcentaje=porcentaje,
        email=email,
        firmas_existentes=json.dumps(firmas_existentes)
    )

@app.route('/descargar/<tipo>/<numero_contrato>')
def descargar(tipo, numero_contrato):
    file_path = os.path.join(OUTPUT_DIR, f"contrato_{numero_contrato}.pdf")
    if not os.path.exists(file_path):
        return f"Error: El archivo {file_path} no existe.", 404
    return send_file(file_path, as_attachment=True)

@app.route('/test_doc')
def test_doc():
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph('Documento de prueba')
        test_path = os.path.join(OUTPUT_DIR, "test.docx")
        doc.save(test_path)
        return f"Documento creado exitosamente en {test_path}", 200
    except Exception as e:
        return f"Error al crear documento: {str(e)}", 500

# Función para limpiar directorios (modificada para scheduler)
def limpiar_directorios():
    """Elimina todos los archivos en OUTPUT_DIR, SIGNATURE_DIR y FIRMAS_TEMP_DIR."""
    try:
        if os.path.exists(OUTPUT_DIR):
            for filename in os.listdir(OUTPUT_DIR):
                file_path = os.path.join(OUTPUT_DIR, filename)
                try:
                    if os.path.isfile(file_path) or os.path.islink(file_path):
                        os.unlink(file_path)
                    elif os.path.isdir(file_path):
                        shutil.rmtree(file_path)
                except Exception as e:
                    print(f"Error al eliminar {file_path}: {e}")
            print(f"Directorio {OUTPUT_DIR} limpiado.")
        
        if os.path.exists(SIGNATURE_DIR):
            for filename in os.listdir(SIGNATURE_DIR):
                file_path = os.path.join(SIGNATURE_DIR, filename)
                try:
                    if os.path.isfile(file_path) or os.path.islink(file_path):
                        os.unlink(file_path)
                    elif os.path.isdir(file_path):
                        shutil.rmtree(file_path)
                except Exception as e:
                    print(f"Error al eliminar {file_path}: {e}")
            print(f"Directorio {SIGNATURE_DIR} limpiado.")
        
        if os.path.exists(FIRMAS_TEMP_DIR):
            for filename in os.listdir(FIRMAS_TEMP_DIR):
                file_path = os.path.join(FIRMAS_TEMP_DIR, filename)
                try:
                    if os.path.isfile(file_path) or os.path.islink(file_path):
                        os.unlink(file_path)
                    elif os.path.isdir(file_path):
                        shutil.rmtree(file_path)
                except Exception as e:
                    print(f"Error al eliminar {file_path}: {e}")
            print(f"Directorio {FIRMAS_TEMP_DIR} limpiado.")
    except Exception as e:
        print(f"Error al limpiar directorios: {e}")

# Configurar la limpieza programada para cada domingo a las 00:00 UTC
def configurar_limpieza_programada():
    scheduler.add_job(
        id='limpieza_semanal',
        func=limpiar_directorios,
        trigger='cron',
        day_of_week='sun',  # Domingo
        hour=0,             # 00:00 (medianoche)
        minute=0,
        second=0,
        timezone='UTC'      # Cambiar a 'America/Bogota' si prefieres hora local
    )
    print("✅ Limpieza programada para cada domingo a las 00:00 UTC")

if __name__ == '__main__':
    import os
    from dotenv import load_dotenv

    # Cargar variables de entorno
    load_dotenv()

    # Inicializar el scheduler con la aplicación Flask
    scheduler.init_app(app)
    
    # Configurar la limpieza programada
    configurar_limpieza_programada()
    
    # Iniciar el scheduler
    scheduler.start()

    # Configuración para Render o desarrollo local
    host = '0.0.0.0'
    port = int(os.getenv('PORT', 5000))
    app.run(host=host, port=port, debug=os.getenv('FLASK_DEBUG', 'False') == 'True')